from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from django.http import HttpResponse
from .forms import TemplateForm
from .models import Template, Draft
from django.template.loader import get_template
from xhtml2pdf import pisa
from docx import Document

# Mock AI 함수 (FastAPI로 교체 예정)
def mock_ai_edit(draft_text, prompt):
    return f"[Edited by AI based on: '{prompt}']\n\n{draft_text}"

def mock_ai_evaluate(draft_text):
    return "🧠 AI 평가 결과: 이 초안은 기술적 진보성이 우수합니다."

def editor_view(request):
    if request.method == 'POST':
        # 템플릿 → 초안 생성
        if 'submit_template' in request.POST:
            form = TemplateForm(request.POST)
            if form.is_valid():
                template = form.save(commit=False)
                # 이 부분 로그인한 사용자 대체
                if request.user.is_authenticated:
                    template.user_code = request.user
                template.date = timezone.now()
                template.save()

                # Draft 생성
                draft = Draft.objects.create(
                    template_id=template,
                    draft_name=f"{template.tech_name}_draft",
                    create_draft=f"[Initial draft based on {template.tech_name}]",
                    create_time=timezone.now()
                )
                return redirect('editor_view')  # 새로고침 후 반영

        # 직접 수정 저장
        elif 'save_draft' in request.POST:
            draft_id = request.POST.get('draft_id')
            new_text = request.POST.get('draft_text')
            draft = Draft.objects.get(draft_id=draft_id)
            draft.create_draft = new_text
            draft.save()

        # AI 수정
        elif 'ai_edit' in request.POST:
            draft_id = request.POST.get('draft_id')
            prompt = request.POST.get('prompt')
            draft = Draft.objects.get(draft_id=draft_id)
            draft.create_draft = mock_ai_edit(draft.create_draft, prompt)
            draft.save()

        # AI 평가
        elif 'ai_evaluate' in request.POST:
            draft_id = request.POST.get('draft_id')
            draft = Draft.objects.get(draft_id=draft_id)
            request.session['ai_feedback'] = mock_ai_evaluate(draft.create_draft)
            request.session['active_draft'] = draft.draft_id

        return redirect('editor_view')

    # GET
    template = Template.objects.last()
    draft = Draft.objects.last()
    feedback = request.session.pop('ai_feedback', None) if draft else None
    return render(request, 'app/editor.html', {
        'form': TemplateForm(),
        'template': template,
        'draft': draft,
        'feedback': feedback
    })


# PDF 및 DOCX 다운로드
def download_pdf(request, draft_id):
    draft = get_object_or_404(Draft, draft_id=draft_id)
    template = get_template('app/draft_pdf_template.html')
    html = template.render({'draft': draft})
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{draft.draft_name}.pdf"'
    pisa.CreatePDF(html, dest=response)
    return response

def download_docx(request, draft_id):
    draft = get_object_or_404(Draft, draft_id=draft_id)
    document = Document()
    document.add_heading(draft.draft_name, level=1)
    document.add_paragraph(draft.create_draft)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{draft.draft_name}.docx"'
    document.save(response)
    return response